"""Convert report.md to report.docx with a clean custom template.

Post-processes the docx to:
- Style figure captions (centered, smaller font)
- Resize images to consistent widths (10cm single-panel, 18cm multi-panel)
- Center all images
- Use A4 paper

Usage:
    uv run --with pypandoc --with pypandoc_binary --with python-docx python misc/convert_md_to_docx.py
"""

import io
import os
import tempfile
import zipfile

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image

import pypandoc

# Target widths in docx (images are generated at 2x these sizes)
WIDTH_SINGLE = Cm(10)
WIDTH_MULTI = Cm(15)

# Images generated at FIG_W1 (~7.9in at 150dpi = ~1181px) are single-panel,
# at FIG_W2 (~14.2in at 150dpi = ~2125px) are multi-panel.
# Mid-point at 1600px to distinguish.
PIXEL_WIDTH_THRESHOLD = 1600


def create_reference_docx(path):
    """Create a clean reference template for pandoc docx conversion."""
    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing = 1.15

    # A4 paper with narrow margins
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Title
    title = doc.styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    title.paragraph_format.space_after = Pt(4)

    # Heading 1
    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(8)

    # Heading 2
    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(6)

    # Heading 3
    h3 = doc.styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    h3.paragraph_format.space_before = Pt(14)
    h3.paragraph_format.space_after = Pt(4)

    # Table
    doc.styles["Table Grid"].font.name = "Calibri"
    doc.styles["Table Grid"].font.size = Pt(10)

    doc.save(path)


def get_image_pixel_widths(docx_path):
    """Read the pixel width of each embedded image from the docx zip, keyed by filename."""
    widths = {}
    with zipfile.ZipFile(docx_path, "r") as z:
        for name in z.namelist():
            if name.startswith("word/media/"):
                data = z.read(name)
                img = Image.open(io.BytesIO(data))
                # Key by just the filename part (e.g., "rId9.png")
                widths[os.path.basename(name)] = img.size[0]
    return widths


def post_process(doc, docx_path):
    """Resize and center images, style figure captions."""
    pixel_widths = get_image_pixel_widths(docx_path)

    # Build rId -> media filename mapping from relationships
    rels = doc.part.rels
    rid_to_file = {}
    for rid, rel in rels.items():
        if "media/" in rel.target_ref:
            rid_to_file[rid] = os.path.basename(rel.target_ref)

    for para in doc.paragraphs:
        # Find inline images in this paragraph
        drawings = para._element.findall(
            ".//" + qn("wp:inline")
        )
        if drawings:
            # Center the paragraph containing the image
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            for drawing in drawings:
                extent = drawing.find(qn("wp:extent"))
                if extent is None:
                    continue

                w_emu = int(extent.get("cx", 0))
                h_emu = int(extent.get("cy", 0))
                if h_emu == 0:
                    continue

                # Find the relationship ID to look up pixel width
                blip = drawing.find(".//" + qn("a:blip"))
                rid = blip.get(qn("r:embed")) if blip is not None else None
                media_file = rid_to_file.get(rid, "")
                px_w = pixel_widths.get(media_file, 0)
                target_w = WIDTH_MULTI if px_w > PIXEL_WIDTH_THRESHOLD else WIDTH_SINGLE

                scale = target_w / w_emu
                new_h = int(h_emu * scale)

                extent.set("cx", str(int(target_w)))
                extent.set("cy", str(new_h))

                # Also update the embedded graphic extent
                for graphic_ext in drawing.findall(".//" + qn("a:ext")):
                    graphic_ext.set("cx", str(int(target_w)))
                    graphic_ext.set("cy", str(new_h))

        # Style figure captions
        if para.text.strip().startswith("Figure "):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.size = Pt(9)
                run.font.italic = True
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Style tables: light gray thin borders, centered
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _set_table_borders(table, color="BBBBBB", size=4)


def _set_table_borders(table, color="BBBBBB", size=4):
    """Set light gray thin borders on all table cells."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")

    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)

    # Remove existing borders if any
    for existing in tbl_pr.findall(qn("w:tblBorders")):
        tbl_pr.remove(existing)
    tbl_pr.append(borders)


if __name__ == "__main__":
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        ref_path = f.name

    try:
        create_reference_docx(ref_path)
        pypandoc.convert_file(
            "report.md",
            "docx",
            outputfile="report.docx",
            extra_args=["--reference-doc=" + ref_path, "--resource-path=."],
        )

        # Post-process
        doc = Document("report.docx")
        post_process(doc, "report.docx")
        doc.save("report.docx")

        print("Saved report.docx")
    finally:
        os.unlink(ref_path)
